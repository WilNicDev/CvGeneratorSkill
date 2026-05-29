#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx>=1.1.0"]
# ///
"""Generate a styled CV Word document from JSON data.

Usage:
    uv run generate_docx.py <input.json> <output.docx>
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_paragraph_spacing(paragraph, before=None, after=None, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))
    if line is not None:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')


def add_bottom_border(paragraph, sz=12, color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_horizontal_rule(doc, color="BFBFBF", sz=4):
    """Add a thin horizontal line as a visual separator between experiences."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=200, after=80, line=240)
    add_bottom_border(p, sz=sz, color=color)
    return p


def add_run(paragraph, text, font_name="Cambria", font_size=Pt(11),
            bold=None, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def keep_with_next(paragraph):
    """Prevent page break between this paragraph and the next."""
    pPr = paragraph._p.get_or_add_pPr()
    kwn = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
    pPr.append(kwn)


def keep_lines_together(paragraph):
    """Prevent page break within this paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    klt = parse_xml(f'<w:keepLines {nsdecls("w")} w:val="true"/>')
    pPr.append(klt)


def add_section_heading(doc, text, space_before=0):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=space_before, after=20, line=240)
    add_bottom_border(p)
    add_run(p, text, font_name="Cambria", font_size=Pt(14),
            bold=True, color=DARK_BLUE)
    keep_with_next(p)
    return p


def add_bullet(doc, text):
    try:
        bp = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(bp, after=0, line=240)
        if bp.runs:
            bp.runs[0].text = text
            bp.runs[0].font.name = "Cambria"
            bp.runs[0].font.size = Pt(11)
        else:
            add_run(bp, text)
        return bp
    except KeyError:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720"/>')
        pPr.append(ind)
        set_paragraph_spacing(p, after=0, line=240)
        add_run(p, text)
        return p


def set_page_margins(section, top=Cm(1), bottom=Cm(1),
                     left=Cm(1.27), right=Cm(1.75)):
    section.top_margin = top
    section.bottom_margin = bottom
    section.left_margin = left
    section.right_margin = right


# Page margin constants in twips for use in raw XML (sectPr).
# python-docx Cm() returns EMU; Word XML pgMar expects twips.
_TWIPS_TOP = 567       # 1 cm
_TWIPS_BOTTOM = 567    # 1 cm
_TWIPS_LEFT = 720      # 1.27 cm
_TWIPS_RIGHT = 992     # 1.75 cm
_TWIPS_PG_W = 11906    # A4 width
_TWIPS_PG_H = 16838    # A4 height


def build_cv(data, output_path):
    doc = Document()

    section = doc.sections[0]
    set_page_margins(section)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

    if doc.paragraphs:
        p = doc.paragraphs[0]._p
        p.getparent().remove(p)

    # 1. Candidate Name (plain paragraph, no banner/shading)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0, line=240)
    add_run(p, data["name"], font_name="Cambria", font_size=Pt(22),
            bold=True, color=DARK_BLUE)

    # 2. Target Headline
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=40, line=240)
    add_run(p, data.get("target_headline", data.get("title", "")),
            font_name="Cambria", font_size=Pt(12), bold=True, color=DARK_BLUE)

    # 3. Contact Line (city/country | phone | email | linkedin)
    contact = data["contact"]
    linkedin_display = contact.get("linkedin", "").replace("https://www.", "").replace("https://", "").replace("http://", "")
    city_country = contact.get("city_country", contact.get("address", ""))
    contact_parts = [city_country, contact["phone"], contact["email"], linkedin_display]
    contact_text = " | ".join(part for part in contact_parts if part)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=40, line=240)
    add_run(p, contact_text, font_name="Cambria", font_size=Pt(10))

    # 4. Languages Line (compact, near top — optional)
    has_languages_line = bool(data.get("languages_line"))
    if has_languages_line:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=120, line=240)
        add_run(p, "Languages: ", font_name="Cambria", font_size=Pt(10),
                bold=True)
        add_run(p, data["languages_line"], font_name="Cambria",
                font_size=Pt(10))

    # 5. Professional Summary
    add_section_heading(doc, "Professional Summary")
    summary = data["profile_summary"]
    if isinstance(summary, str):
        summary = [summary]
    for para_text in summary:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=60, line=240)
        run = p.add_run(para_text)
        run.font.name = "Cambria"
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK

    # 6. Core Expertise (comma-separated list for ATS compatibility)
    add_section_heading(doc, "Core Expertise", space_before=360)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0, line=240)
    add_run(p, ", ".join(data["key_skills"]))

    # 7. Career Highlights
    add_section_heading(doc, "Career Highlights", space_before=360)
    for hl in data["career_highlights"]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=60, line=240)
        add_run(p, f"{hl['company']}: ", bold=True)
        add_run(p, hl["text"])

    # 8. Professional Experience
    add_section_heading(doc, "Professional Experience", space_before=360)
    for i, exp in enumerate(data["experiences"]):
        if i > 0:
            add_horizontal_rule(doc)
        role_line = f"{exp['date_range']}: {exp['company']} – {exp['role']} ({exp['employment_type']})"
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=160 if i == 0 else 40, after=60, line=240)
        add_run(p, role_line, bold=True)
        keep_with_next(p)
        if exp.get("company_context"):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=60, line=240)
            add_run(p, exp["company_context"])
            keep_with_next(p)
        if exp.get("responsibilities"):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=60, after=20, line=240)
            add_run(p, "Responsibilities:", bold=True)
            keep_with_next(p)
            for resp in exp["responsibilities"]:
                add_bullet(doc, resp)
        if exp.get("achievements"):
            heading = exp.get("achievements_heading", "Key Achievements:")
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=60, after=20, line=240)
            add_run(p, heading, bold=True)
            keep_with_next(p)
            for ach in exp["achievements"]:
                if exp.get("use_bullet_achievements", True):
                    add_bullet(doc, ach)
                else:
                    p = doc.add_paragraph()
                    set_paragraph_spacing(p, after=40, line=240)
                    add_run(p, ach)
        if exp.get("additional_skills"):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=120, after=0, line=240)
            add_run(p, "Additional skills: ", bold=True)
            add_run(p, ", ".join(exp["additional_skills"]))
        if exp.get("platforms_and_tools"):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=0, line=240)
            add_run(p, "Platforms & tools: ", bold=True)
            add_run(p, ", ".join(exp["platforms_and_tools"]))

    # 9. Earlier Career
    if data.get("early_career"):
        add_section_heading(doc, "Earlier Career", space_before=360)
        for ec in data["early_career"]:
            date_info = ec.get("date_range", "")
            line = f"{date_info}: {ec['company']} – {ec['role']} ({ec['employment_type']})"
            if ec.get("note"):
                line += f" — {ec['note']}"
            add_bullet(doc, line)

    # 10. Education / Qualifications & Certifications
    add_section_heading(doc, "Education / Qualifications", space_before=360)
    for edu in data["education"]:
        add_bullet(doc, f"{edu['programme']}: {edu['institution']} ({edu['year']})")
    if data.get("certifications"):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=60, after=0, line=240)
        add_run(p, "Certifications:", bold=True)
        keep_with_next(p)
        for cert in data["certifications"]:
            label = cert["certification"]
            if cert.get("expired"):
                label += " (expired)"
            add_bullet(doc, f"{label}: {cert['issuer']}")

    # Languages (standalone section — only if no languages_line near top)
    if not has_languages_line:
        add_section_heading(doc, "Languages", space_before=360)
        for lang in data["languages"]:
            line = f"{lang['language']}: {lang['proficiency']}"
            add_bullet(doc, line)

    # 11. Additional Skills & Tools
    appendix = data.get("additional_skills_tools")
    if appendix:
        has_skills = appendix.get("skills") and len(appendix["skills"]) > 0
        has_tools = appendix.get("tools_and_platforms") and len(appendix["tools_and_platforms"]) > 0
        if has_skills or has_tools:
            add_section_heading(doc, "Additional Skills & Tools", space_before=360)
            if has_skills:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, after=120, line=276)
                add_run(p, "Skills: ", bold=True)
                add_run(p, ", ".join(appendix["skills"]))
            if has_tools:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, after=120, line=276)
                add_run(p, "Tools & Platforms: ", bold=True)
                add_run(p, ", ".join(appendix["tools_and_platforms"]))

    doc.save(output_path)
    print(f"CV saved to {output_path}")


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 generate_docx.py <input.json> <output.docx>",
              file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    build_cv(data, sys.argv[2])


if __name__ == "__main__":
    main()
